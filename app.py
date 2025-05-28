from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify
from flask_sqlalchemy import SQLAlchemy
import markdown
import os
from datetime import datetime
import tempfile
import pymysql
from config import config
import zipfile
from docx import Document
from docx.shared import Inches
import re
import json
from functools import wraps
import requests
from urllib.parse import urlparse
import io
from PIL import Image

# 安装PyMySQL作为MySQLdb的替代
pymysql.install_as_MySQLdb()

# 创建Flask应用
def create_app(config_name=None):
    """应用工厂函数"""
    app = Flask(__name__)
    
    # 加载配置
    config_name = config_name or os.environ.get('FLASK_ENV', 'development')
    app.config.from_object(config[config_name])
    
    return app

app = create_app()
db = SQLAlchemy(app)

# 数据库模型
class MarkdownDocument(db.Model):
    __tablename__ = 'markdown_documents'
    
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(255), nullable=False)
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    def __repr__(self):
        return f'<MarkdownDocument {self.title}>'
    
    def to_dict(self):
        """转换为字典格式"""
        return {
            'id': self.id,
            'title': self.title,
            'content': self.content,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

# 辅助函数
def require_api_key(f):
    """API密钥验证装饰器"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # 如果配置中禁用了API密钥验证，直接通过
        if not app.config.get('REQUIRE_API_KEY', True):
            return f(*args, **kwargs)
        
        # 获取API密钥
        api_key = request.headers.get(app.config.get('API_KEY_HEADER', 'X-API-Key'))
        
        # 验证API密钥
        if not api_key:
            return jsonify({
                'success': False,
                'error': f'缺少API密钥，请在请求头中添加 {app.config.get("API_KEY_HEADER", "X-API-Key")}'
            }), 401
        
        if api_key != app.config.get('API_KEY'):
            return jsonify({
                'success': False,
                'error': 'API密钥无效'
            }), 401
        
        return f(*args, **kwargs)
    return decorated_function

def extract_title_from_content(content):
    """从Markdown内容中自动提取标题"""
    if not content or not content.strip():
        return "无标题文档"
    
    lines = content.strip().split('\n')
    
    # 优先查找一级标题 (# 标题)
    for line in lines:
        line = line.strip()
        if line.startswith('# '):
            title = line[2:].strip()
            if title:
                return title
    
    # 查找二级标题 (## 标题)
    for line in lines:
        line = line.strip()
        if line.startswith('## '):
            title = line[3:].strip()
            if title:
                return title
    
    # 查找任意级别标题
    for line in lines:
        line = line.strip()
        if line.startswith('#'):
            title = line.lstrip('#').strip()
            if title:
                return title
    
    # 如果没有找到标题，使用第一行非空内容（限制长度）
    for line in lines:
        line = line.strip()
        if line and not line.startswith('#'):
            # 移除Markdown格式符号
            clean_line = re.sub(r'[*_`~\[\]()]+', '', line)
            clean_line = clean_line.strip()
            if clean_line:
                # 限制标题长度
                return clean_line[:50] + ('...' if len(clean_line) > 50 else '')
    
    return "无标题文档"

def download_image(url, max_size=(800, 600)):
    """下载图片并调整大小"""
    try:
        # 检查URL是否有效
        parsed_url = urlparse(url)
        if not parsed_url.scheme or not parsed_url.netloc:
            return None
        
        # 设置请求头，模拟浏览器访问
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        # 下载图片，设置超时
        response = requests.get(url, headers=headers, timeout=10, stream=True)
        response.raise_for_status()
        
        # 检查内容类型
        content_type = response.headers.get('content-type', '').lower()
        if not content_type.startswith('image/'):
            return None
        
        # 读取图片数据
        image_data = io.BytesIO()
        for chunk in response.iter_content(chunk_size=8192):
            image_data.write(chunk)
        image_data.seek(0)
        
        # 使用PIL处理图片
        try:
            with Image.open(image_data) as img:
                # 转换为RGB模式（如果是RGBA或其他模式）
                if img.mode in ('RGBA', 'LA', 'P'):
                    # 创建白色背景
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # 调整图片大小，保持宽高比
                img.thumbnail(max_size, Image.Resampling.LANCZOS)
                
                # 保存到内存
                output = io.BytesIO()
                img.save(output, format='JPEG', quality=85, optimize=True)
                output.seek(0)
                
                return output
        except Exception as e:
            print(f"图片处理失败: {e}")
            return None
            
    except Exception as e:
        print(f"下载图片失败 {url}: {e}")
        return None

def render_markdown(content):
    """渲染Markdown内容为HTML"""
    return markdown.markdown(
        content, 
        extensions=app.config['MARKDOWN_EXTENSIONS']
    )

def create_download_file(document):
    """创建下载文件"""
    content = f"""# {document.title}

**创建时间**: {document.created_at.strftime('%Y-%m-%d %H:%M:%S')}  
**更新时间**: {document.updated_at.strftime('%Y-%m-%d %H:%M:%S')}

---

{document.content}
"""
    
    # 创建临时文件
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_file:
        temp_file.write(content)
        return temp_file.name

def markdown_to_word(document):
    """将Markdown文档转换为Word文档，支持图片嵌入"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading(document.title, 0)
    
    # 添加元数据
    meta_para = doc.add_paragraph()
    meta_para.add_run('创建时间: ').bold = True
    meta_para.add_run(document.created_at.strftime('%Y-%m-%d %H:%M:%S'))
    meta_para.add_run('\n更新时间: ').bold = True
    meta_para.add_run(document.updated_at.strftime('%Y-%m-%d %H:%M:%S'))
    
    # 添加分隔线
    doc.add_paragraph('─' * 50)
    
    # 处理Markdown内容
    lines = document.content.split('\n')
    current_para = None
    in_code_block = False
    code_lang = None
    
    for line in lines:
        line = line.rstrip()
        
        # 处理代码块
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line[3:].strip() if len(line) > 3 else ''
                current_para = doc.add_paragraph()
                current_para.style = 'No Spacing'
                continue
            else:
                in_code_block = False
                current_para = None
                continue
        
        if in_code_block:
            if current_para is None:
                current_para = doc.add_paragraph()
                current_para.style = 'No Spacing'
            run = current_para.add_run(line + '\n')
            run.font.name = 'Courier New'
            continue
        
        # 处理标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()
            if level <= 6 and title_text:
                doc.add_heading(title_text, level)
                continue
        
        # 处理列表
        if line.startswith('- ') or line.startswith('* '):
            # 检查列表项中是否有图片
            list_text = line[2:].strip()
            if process_line_with_images(doc, list_text, style='List Bullet'):
                continue
            else:
                para = doc.add_paragraph(list_text, style='List Bullet')
                continue
        
        if re.match(r'^\d+\.\s', line):
            list_text = re.sub(r'^\d+\.\s', '', line)
            if process_line_with_images(doc, list_text, style='List Number'):
                continue
            else:
                para = doc.add_paragraph(list_text, style='List Number')
                continue
        
        # 处理引用
        if line.startswith('>'):
            quote_text = line[1:].strip()
            if process_line_with_images(doc, quote_text, style='Quote'):
                continue
            else:
                para = doc.add_paragraph(quote_text)
                para.style = 'Quote'
                continue
        
        # 处理图片（独立行）
        img_match = re.match(r'^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$', line)
        if img_match:
            alt_text = img_match.group(1)
            img_url = img_match.group(2)
            
            # 下载并插入图片
            image_data = download_image(img_url)
            if image_data:
                try:
                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(image_data, width=Inches(5))  # 设置图片宽度为5英寸
                    
                    # 如果有alt文本，添加图片说明
                    if alt_text:
                        caption_para = doc.add_paragraph()
                        caption_run = caption_para.add_run(f"图片: {alt_text}")
                        caption_run.italic = True
                        caption_para.alignment = 1  # 居中对齐
                except Exception as e:
                    print(f"插入图片失败: {e}")
                    # 如果插入失败，显示图片链接
                    para = doc.add_paragraph()
                    run = para.add_run(f"[图片: {alt_text}]({img_url})")
                    run.italic = True
            else:
                # 如果下载失败，显示图片链接
                para = doc.add_paragraph()
                run = para.add_run(f"[图片: {alt_text}]({img_url})")
                run.italic = True
            continue
        
        # 处理普通段落
        if line.strip():
            if not process_line_with_images(doc, line):
                # 如果没有图片，按原来的方式处理
                para = doc.add_paragraph()
                process_inline_formatting(para, line)
        else:
            # 空行
            doc.add_paragraph()
    
    # 保存到临时文件
    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(temp_file.name)
    return temp_file.name

def process_line_with_images(doc, text, style=None):
    """处理包含图片的文本行"""
    # 查找图片模式 ![alt](url)
    img_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    
    if not re.search(img_pattern, text):
        return False  # 没有图片
    
    # 创建段落
    para = doc.add_paragraph()
    if style:
        para.style = style
    
    # 分割文本和图片
    parts = re.split(img_pattern, text)
    
    i = 0
    while i < len(parts):
        if i % 3 == 0:
            # 普通文本
            if parts[i]:
                process_inline_formatting(para, parts[i])
        elif i % 3 == 1:
            # alt文本
            alt_text = parts[i]
        elif i % 3 == 2:
            # 图片URL
            img_url = parts[i]
            
            # 下载并插入图片
            image_data = download_image(img_url)
            if image_data:
                try:
                    run = para.add_run()
                    run.add_picture(image_data, width=Inches(3))  # 行内图片使用较小尺寸
                except Exception as e:
                    print(f"插入行内图片失败: {e}")
                    # 如果插入失败，显示图片链接
                    run = para.add_run(f"[图片: {alt_text}]({img_url})")
                    run.italic = True
            else:
                # 如果下载失败，显示图片链接
                run = para.add_run(f"[图片: {alt_text}]({img_url})")
                run.italic = True
        
        i += 1
    
    return True

def process_inline_formatting(para, text):
    """处理内联格式（粗体、斜体等）"""
    parts = []
    
    # 处理粗体 **text**
    while '**' in text:
        start = text.find('**')
        if start == -1:
            break
        end = text.find('**', start + 2)
        if end == -1:
            break
        
        if start > 0:
            parts.append(('normal', text[:start]))
        parts.append(('bold', text[start+2:end]))
        text = text[end+2:]
    
    if text:
        parts.append(('normal', text))
    
    # 添加格式化文本
    for part_type, part_text in parts:
        run = para.add_run(part_text)
        if part_type == 'bold':
            run.bold = True

def create_batch_zip(documents, file_type='md'):
    """创建批量下载的ZIP文件"""
    temp_zip = tempfile.NamedTemporaryFile(suffix='.zip', delete=False)
    
    with zipfile.ZipFile(temp_zip.name, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for doc in documents:
            if file_type == 'md':
                # 创建Markdown文件
                temp_file = create_download_file(doc)
                safe_title = "".join(c for c in doc.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
                # 确保文件名不为空
                if not safe_title:
                    safe_title = f"document_{doc.id}"
                filename = f"{safe_title}.md"
                zipf.write(temp_file, filename.encode('utf-8').decode('utf-8'))
                os.remove(temp_file)
            elif file_type == 'docx':
                # 创建Word文件
                temp_file = markdown_to_word(doc)
                safe_title = "".join(c for c in doc.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
                # 确保文件名不为空
                if not safe_title:
                    safe_title = f"document_{doc.id}"
                filename = f"{safe_title}.docx"
                zipf.write(temp_file, filename.encode('utf-8').decode('utf-8'))
                os.remove(temp_file)
    
    return temp_zip.name

# 路由
@app.route('/')
def index():
    """首页 - 显示所有Markdown文档列表"""
    try:
        page = request.args.get('page', 1, type=int)
        per_page = 10  # 每页显示10个文档
        
        documents = MarkdownDocument.query.order_by(MarkdownDocument.created_at.desc()).paginate(
            page=page, per_page=per_page, error_out=False
        )
        return render_template('index.html', documents=documents)
    except Exception as e:
        flash(f'获取文档列表失败: {str(e)}', 'error')
        return render_template('index.html', documents=None)

@app.route('/create', methods=['GET', 'POST'])
def create_document():
    """创建新的Markdown文档"""
    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        content = request.form.get('content', '').strip()
        
        # 验证输入
        if not title:
            flash('文档标题不能为空', 'error')
            return render_template('create.html', title=title, content=content)
        
        if not content:
            flash('文档内容不能为空', 'error')
            return render_template('create.html', title=title, content=content)
        
        try:
            document = MarkdownDocument(title=title, content=content)
            db.session.add(document)
            db.session.commit()
            
            flash('文档创建成功！', 'success')
            return redirect(url_for('view_document', doc_id=document.id))
        except Exception as e:
            db.session.rollback()
            flash(f'创建文档失败: {str(e)}', 'error')
            return render_template('create.html', title=title, content=content)
    
    return render_template('create.html')

@app.route('/view/<int:doc_id>')
def view_document(doc_id):
    """查看Markdown文档（渲染后的HTML）"""
    try:
        document = MarkdownDocument.query.get_or_404(doc_id)
        html_content = render_markdown(document.content)
        return render_template('view.html', document=document, html_content=html_content)
    except Exception as e:
        flash(f'获取文档失败: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/edit/<int:doc_id>', methods=['GET', 'POST'])
def edit_document(doc_id):
    """编辑Markdown文档"""
    try:
        document = MarkdownDocument.query.get_or_404(doc_id)
        
        if request.method == 'POST':
            title = request.form.get('title', '').strip()
            content = request.form.get('content', '').strip()
            
            # 验证输入
            if not title:
                flash('文档标题不能为空', 'error')
                return render_template('edit.html', document=document)
            
            if not content:
                flash('文档内容不能为空', 'error')
                return render_template('edit.html', document=document)
            
            try:
                document.title = title
                document.content = content
                document.updated_at = datetime.utcnow()
                
                db.session.commit()
                flash('文档更新成功！', 'success')
                return redirect(url_for('view_document', doc_id=doc_id))
            except Exception as e:
                db.session.rollback()
                flash(f'更新文档失败: {str(e)}', 'error')
        
        return render_template('edit.html', document=document)
    except Exception as e:
        flash(f'获取文档失败: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/delete/<int:doc_id>')
def delete_document(doc_id):
    """删除Markdown文档"""
    try:
        document = MarkdownDocument.query.get_or_404(doc_id)
        db.session.delete(document)
        db.session.commit()
        flash('文档删除成功！', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'删除文档失败: {str(e)}', 'error')
    
    return redirect(url_for('index'))

@app.route('/download/<int:doc_id>')
def download_document(doc_id):
    """下载Markdown文档"""
    try:
        document = MarkdownDocument.query.get_or_404(doc_id)
        
        # 创建下载文件
        temp_filename = create_download_file(document)
        
        # 生成安全的文件名
        safe_title = "".join(c for c in document.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        download_filename = f"{safe_title}.md"
        
        def remove_file(response):
            """响应后删除临时文件"""
            try:
                os.remove(temp_filename)
            except Exception:
                pass
            return response
        
        return send_file(
            temp_filename,
            as_attachment=True,
            download_name=download_filename,
            mimetype='text/markdown'
        )
    except Exception as e:
        flash(f'下载文档失败: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/download_word/<int:doc_id>')
def download_word(doc_id):
    """下载Word文档"""
    try:
        document = MarkdownDocument.query.get_or_404(doc_id)
        
        # 创建Word文件
        temp_filename = markdown_to_word(document)
        
        # 生成安全的文件名
        safe_title = "".join(c for c in document.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        download_filename = f"{safe_title}.docx"
        
        def remove_file(response):
            """响应后删除临时文件"""
            try:
                os.remove(temp_filename)
            except Exception:
                pass
            return response
        
        return send_file(
            temp_filename,
            as_attachment=True,
            download_name=download_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        flash(f'下载Word文档失败: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/batch_download', methods=['POST'])
def batch_download():
    """批量下载文档"""
    try:
        data = request.get_json()
        doc_ids = data.get('doc_ids', [])
        file_type = data.get('file_type', 'md')  # 'md' 或 'docx'
        
        if not doc_ids:
            return jsonify({'error': '请选择要下载的文档'}), 400
        
        # 获取文档
        documents = MarkdownDocument.query.filter(MarkdownDocument.id.in_(doc_ids)).all()
        
        if not documents:
            return jsonify({'error': '未找到指定的文档'}), 404
        
        # 创建ZIP文件
        zip_filename = create_batch_zip(documents, file_type)
        
        # 生成下载文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        download_filename = f"markdown_documents_{timestamp}.zip"
        
        return send_file(
            zip_filename,
            as_attachment=True,
            download_name=download_filename,
            mimetype='application/zip'
        )
    except Exception as e:
        return jsonify({'error': f'批量下载失败: {str(e)}'}), 500

@app.route('/batch_delete', methods=['POST'])
def batch_delete():
    """批量删除文档"""
    try:
        data = request.get_json()
        doc_ids = data.get('doc_ids', [])
        
        if not doc_ids:
            return jsonify({'error': '请选择要删除的文档'}), 400
        
        # 删除文档
        deleted_count = MarkdownDocument.query.filter(MarkdownDocument.id.in_(doc_ids)).delete()
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'成功删除 {deleted_count} 个文档'
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'批量删除失败: {str(e)}'}), 500

# API接口
@app.route('/api/documents', methods=['POST'])
@require_api_key
def api_create_document():
    """API接口：创建文档
    
    请求格式:
    {
        "content": "Markdown内容"
    }
    
    响应格式:
    {
        "success": true,
        "data": {
            "id": 1,
            "title": "自动提取的标题",
            "content": "文档内容",
            "created_at": "2023-01-01T00:00:00",
            "updated_at": "2023-01-01T00:00:00"
        },
        "message": "文档创建成功"
    }
    """
    try:
        # 检查Content-Type
        if not request.is_json:
            return jsonify({
                'success': False,
                'error': 'Content-Type必须是application/json'
            }), 400
        
        data = request.get_json()
        
        # 验证必需参数
        if not data or 'content' not in data:
            return jsonify({
                'success': False,
                'error': '缺少必需参数: content'
            }), 400
        
        content = data.get('content', '').strip()
        
        # 验证内容不为空
        if not content:
            return jsonify({
                'success': False,
                'error': '文档内容不能为空'
            }), 400
        
        # 自动提取标题
        title = extract_title_from_content(content)
        
        # 创建文档
        document = MarkdownDocument(
            title=title,
            content=content
        )
        
        db.session.add(document)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'data': document.to_dict(),
            'message': '文档创建成功'
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': f'创建文档失败: {str(e)}'
        }), 500

@app.route('/api/documents', methods=['GET'])
@require_api_key
def api_list_documents():
    """API接口：获取文档列表
    
    查询参数:
    - page: 页码 (默认: 1)
    - per_page: 每页数量 (默认: 10, 最大: 100)
    
    响应格式:
    {
        "success": true,
        "data": {
            "documents": [...],
            "pagination": {
                "page": 1,
                "per_page": 10,
                "total": 100,
                "pages": 10,
                "has_prev": false,
                "has_next": true
            }
        }
    }
    """
    try:
        page = request.args.get('page', 1, type=int)
        per_page = min(request.args.get('per_page', 10, type=int), 100)  # 限制最大每页数量
        
        documents = MarkdownDocument.query.order_by(
            MarkdownDocument.updated_at.desc()
        ).paginate(
            page=page,
            per_page=per_page,
            error_out=False
        )
        
        return jsonify({
            'success': True,
            'data': {
                'documents': [doc.to_dict() for doc in documents.items],
                'pagination': {
                    'page': documents.page,
                    'per_page': documents.per_page,
                    'total': documents.total,
                    'pages': documents.pages,
                    'has_prev': documents.has_prev,
                    'has_next': documents.has_next
                }
            }
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取文档列表失败: {str(e)}'
        }), 500

@app.route('/api/documents/<int:doc_id>', methods=['GET'])
@require_api_key
def api_get_document(doc_id):
    """API接口：获取单个文档
    
    响应格式:
    {
        "success": true,
        "data": {
            "id": 1,
            "title": "文档标题",
            "content": "文档内容",
            "created_at": "2023-01-01T00:00:00",
            "updated_at": "2023-01-01T00:00:00"
        }
    }
    """
    try:
        document = MarkdownDocument.query.get(doc_id)
        
        if not document:
            return jsonify({
                'success': False,
                'error': '文档不存在'
            }), 404
        
        return jsonify({
            'success': True,
            'data': document.to_dict()
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取文档失败: {str(e)}'
        }), 500

@app.route('/api/info', methods=['GET'])
def api_info():
    """API接口：获取API信息（无需验证）
    
    响应格式:
    {
        "success": true,
        "data": {
            "api_version": "1.0.0",
            "require_api_key": true,
            "api_key_header": "X-API-Key",
            "endpoints": [...]
        }
    }
    """
    try:
        return jsonify({
            'success': True,
            'data': {
                'api_version': '1.0.0',
                'require_api_key': app.config.get('REQUIRE_API_KEY', True),
                'api_key_header': app.config.get('API_KEY_HEADER', 'X-API-Key'),
                'endpoints': [
                    {
                        'method': 'POST',
                        'path': '/api/documents',
                        'description': '创建文档',
                        'requires_auth': app.config.get('REQUIRE_API_KEY', True)
                    },
                    {
                        'method': 'GET',
                        'path': '/api/documents',
                        'description': '获取文档列表',
                        'requires_auth': app.config.get('REQUIRE_API_KEY', True)
                    },
                    {
                        'method': 'GET',
                        'path': '/api/documents/{doc_id}',
                        'description': '获取单个文档',
                        'requires_auth': app.config.get('REQUIRE_API_KEY', True)
                    },
                    {
                        'method': 'GET',
                        'path': '/api/info',
                        'description': '获取API信息',
                        'requires_auth': False
                    }
                ]
            }
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取API信息失败: {str(e)}'
        }), 500

# 错误处理
@app.errorhandler(404)
def not_found_error(error):
    """404错误处理"""
    return render_template('base.html'), 404

@app.errorhandler(500)
def internal_error(error):
    """500错误处理"""
    db.session.rollback()
    return render_template('base.html'), 500

# 模板上下文处理器
@app.context_processor
def inject_config():
    """注入配置到模板上下文"""
    return {
        'app_name': 'Markdown Factory',
        'version': '1.0.0'
    }

if __name__ == '__main__':
    with app.app_context():
        try:
            db.create_all()
            print("✅ 数据库表创建成功")
        except Exception as e:
            print(f"❌ 数据库初始化失败: {e}")
    
    app.run(debug=True, host='0.0.0.0', port=8888) 